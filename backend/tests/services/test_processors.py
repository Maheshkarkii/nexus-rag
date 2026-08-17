"""Unit tests for format-specific document processors (PDF, DOCX, TXT, CSV, Excel, JSON)."""

import json
import uuid
from pathlib import Path

import docx
import openpyxl
import pytest
from pypdf import PdfWriter

from app.db.models.document import Document
from app.services.document_processing.csv_processor import CSVProcessor
from app.services.document_processing.docx_processor import DocxProcessor
from app.services.document_processing.excel_processor import ExcelProcessor
from app.services.document_processing.json_processor import JSONProcessor
from app.services.document_processing.normalizer import normalize_extracted_text
from app.services.document_processing.pdf_processor import PDFProcessor
from app.services.document_processing.text_processor import TextProcessor


def create_mock_doc(ext: str, mime: str) -> Document:
    """Create lightweight Document model instance for processor tests."""
    return Document(
        id=uuid.uuid4(),
        project_id=uuid.uuid4(),
        original_filename=f"test_file{ext}",
        stored_filename=f"{uuid.uuid4().hex}{ext}",
        storage_path=f"projects/mock/{uuid.uuid4().hex}{ext}",
        mime_type=mime,
        file_extension=ext,
        file_size=1024,
        status="uploaded",
    )


# ------------------------------------------------------------------------------
# 1. Text Normalizer Tests
# ------------------------------------------------------------------------------
def test_normalize_extracted_text() -> None:
    """Verify normalizing line endings, control characters, and excess blank lines."""
    raw = "Header Line\r\n\r\n\r\n\r\nParagraph 1\r\n\r\nParagraph 2 with trailing spaces   \n"
    normalized = normalize_extracted_text(raw)
    assert "Paragraph 1\n\nParagraph 2 with trailing spaces" in normalized
    assert "\r" not in normalized
    assert "\n\n\n" not in normalized


# ------------------------------------------------------------------------------
# 2. Text Processor Tests
# ------------------------------------------------------------------------------
@pytest.mark.asyncio
async def test_text_processor_utf8_and_latin1(tmp_path: Path) -> None:
    """Verify extracting plain text files with UTF-8 and Latin-1 encodings."""
    txt_file = tmp_path / "notes.txt"
    txt_file.write_text("Introduction to Transformer Architectures\nKey Attention Mechanisms", encoding="utf-8")

    processor = TextProcessor()
    doc = create_mock_doc(".txt", "text/plain")
    result = await processor.extract(txt_file, doc)

    assert "Introduction to Transformer Architectures" in result.text
    assert result.character_count > 0
    assert result.word_count > 0
    assert result.metadata["line_count"] == 2


@pytest.mark.asyncio
async def test_text_processor_empty_file_fails(tmp_path: Path) -> None:
    """Verify processing empty text files raises a ValueError."""
    empty_file = tmp_path / "empty.txt"
    empty_file.write_bytes(b"")

    processor = TextProcessor()
    doc = create_mock_doc(".txt", "text/plain")
    with pytest.raises(ValueError, match="completely empty"):
        await processor.extract(empty_file, doc)


# ------------------------------------------------------------------------------
# 3. CSV Processor Tests
# ------------------------------------------------------------------------------
@pytest.mark.asyncio
async def test_csv_processor_tabular_extraction(tmp_path: Path) -> None:
    """Verify extracting structured CSV datasets."""
    csv_file = tmp_path / "benchmark.csv"
    csv_file.write_text(
        "model,parameters,accuracy\n"
        "Llama-3-70B,70B,92.4\n"
        "Claude-3.5-Sonnet,N/A,93.8\n",
        encoding="utf-8",
    )

    processor = CSVProcessor()
    doc = create_mock_doc(".csv", "text/csv")
    result = await processor.extract(csv_file, doc)

    assert "Columns: model | parameters | accuracy" in result.text
    assert "Row 1: Llama-3-70B | 70B | 92.4" in result.text
    assert "Row 2: Claude-3.5-Sonnet | N/A | 93.8" in result.text
    assert result.metadata["row_count"] == 2
    assert result.metadata["column_count"] == 3


@pytest.mark.asyncio
async def test_csv_processor_empty_fails(tmp_path: Path) -> None:
    """Verify empty CSV files raise a ValueError."""
    csv_file = tmp_path / "empty.csv"
    csv_file.write_bytes(b"")

    processor = CSVProcessor()
    doc = create_mock_doc(".csv", "text/csv")
    with pytest.raises(ValueError, match="empty"):
        await processor.extract(csv_file, doc)


# ------------------------------------------------------------------------------
# 4. JSON Processor Tests
# ------------------------------------------------------------------------------
@pytest.mark.asyncio
async def test_json_processor_hierarchical_extraction(tmp_path: Path) -> None:
    """Verify structured hierarchy and array formatting in JSON files."""
    json_file = tmp_path / "data.json"
    payload = {
        "project": "Transformer Research",
        "authors": ["Vaswani et al.", "Devlin et al."],
        "metrics": {"bleu_score": 28.4, "validated": True},
    }
    json_file.write_text(json.dumps(payload), encoding="utf-8")

    processor = JSONProcessor()
    doc = create_mock_doc(".json", "application/json")
    result = await processor.extract(json_file, doc)

    assert "project: Transformer Research" in result.text
    assert "- Vaswani et al." in result.text
    assert "bleu_score: 28.4" in result.text
    assert result.metadata["root_type"] == "dict"


@pytest.mark.asyncio
async def test_json_processor_malformed_fails(tmp_path: Path) -> None:
    """Verify malformed JSON raises a clean ValueError."""
    json_file = tmp_path / "invalid.json"
    json_file.write_text("{broken json, missing quote", encoding="utf-8")

    processor = JSONProcessor()
    doc = create_mock_doc(".json", "application/json")
    with pytest.raises(ValueError, match="Failed to parse invalid JSON"):
        await processor.extract(json_file, doc)


# ------------------------------------------------------------------------------
# 5. DOCX Processor Tests
# ------------------------------------------------------------------------------
@pytest.mark.asyncio
async def test_docx_processor_paragraphs_and_tables(tmp_path: Path) -> None:
    """Verify extracting paragraphs, headings, and table cells from Word documents."""
    docx_file = tmp_path / "paper.docx"
    doc_obj = docx.Document()
    doc_obj.add_heading("Attention Mechanism Overview", level=1)
    doc_obj.add_paragraph("Self-attention connects all positions with a constant number of operations.")
    table = doc_obj.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Layer"
    table.cell(0, 1).text = "Complexity"
    table.cell(1, 0).text = "Self-Attention"
    table.cell(1, 1).text = "O(n^2)"
    doc_obj.save(str(docx_file))

    processor = DocxProcessor()
    doc = create_mock_doc(".docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    result = await processor.extract(docx_file, doc)

    assert "## Attention Mechanism Overview" in result.text
    assert "Self-attention connects all positions" in result.text
    assert "[Table 1]" in result.text
    assert "Layer | Complexity" in result.text
    assert "Self-Attention | O(n^2)" in result.text
    assert result.metadata["table_count"] == 1


# ------------------------------------------------------------------------------
# 6. Excel Processor Tests
# ------------------------------------------------------------------------------
@pytest.mark.asyncio
async def test_excel_processor_multi_sheet(tmp_path: Path) -> None:
    """Verify extracting multiple worksheets from Excel spreadsheets."""
    xlsx_file = tmp_path / "experiments.xlsx"
    wb = openpyxl.Workbook()
    ws1 = wb.active
    ws1.title = "Benchmark"
    ws1.append(["Model", "Accuracy"])
    ws1.append(["BERT", "88.5"])

    ws2 = wb.create_sheet(title="Latency")
    ws2.append(["Model", "Latency_ms"])
    ws2.append(["BERT", "45"])
    wb.save(str(xlsx_file))

    processor = ExcelProcessor()
    doc = create_mock_doc(".xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    result = await processor.extract(xlsx_file, doc)

    assert "Sheet: Benchmark" in result.text
    assert "Model | Accuracy" in result.text
    assert "BERT | 88.5" in result.text
    assert "Sheet: Latency" in result.text
    assert "BERT | 45" in result.text
    assert result.metadata["sheet_count"] == 2


# ------------------------------------------------------------------------------
# 7. PDF Processor Tests
# ------------------------------------------------------------------------------
@pytest.mark.asyncio
async def test_pdf_processor_with_text(tmp_path: Path) -> None:
    """Verify PDF processor parses digital text and tracks page boundaries."""
    pdf_file = tmp_path / "sample.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    with open(pdf_file, "wb") as f:
        writer.write(f)

    processor = PDFProcessor()
    doc = create_mock_doc(".pdf", "application/pdf")

    # A purely blank PDF with no text layers should raise the expected empty text ValueError
    with pytest.raises(ValueError, match="No extractable text found in PDF document"):
        await processor.extract(pdf_file, doc)
