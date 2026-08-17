"""Integration tests for document text extraction pipeline, content inspection, and error handling."""

import io
import json
import uuid
import docx
from fastapi.testclient import TestClient
import openpyxl
import pytest


def create_sample_project(client: TestClient) -> str:
    """Helper fixture to create a valid research project workspace."""
    res = client.post("/api/v1/projects", json={"name": "Text Processing Workspace"})
    assert res.status_code == 201
    return res.json()["id"]


# ------------------------------------------------------------------------------
# 1. Plain Text Processing Test
# ------------------------------------------------------------------------------
def test_process_txt_document(client: TestClient) -> None:
    """Verify uploading and processing a .txt document extracts normalized text."""
    project_id = create_sample_project(client)
    content = b"Research Note 101:\nSelf-attention allows the model to attend to different representation subspaces."

    upload_res = client.post(
        f"/api/v1/projects/{project_id}/documents",
        files={"file": ("notes.txt", io.BytesIO(content), "text/plain")},
    )
    assert upload_res.status_code == 201
    doc_id = upload_res.json()["id"]
    assert upload_res.json()["status"] == "uploaded"

    # Trigger processing
    proc_res = client.post(f"/api/v1/projects/{project_id}/documents/{doc_id}/process")
    assert proc_res.status_code == 200
    data = proc_res.json()
    assert data["status"] == "extracted"
    assert data["extracted_character_count"] > 0
    assert data["extracted_word_count"] > 0
    assert data["processed_at"] is not None

    # Inspect extracted content
    content_res = client.get(f"/api/v1/projects/{project_id}/documents/{doc_id}/content")
    assert content_res.status_code == 200
    c_data = content_res.json()
    assert "Self-attention allows the model" in c_data["extracted_text"]
    assert c_data["status"] == "extracted"


# ------------------------------------------------------------------------------
# 2. CSV Tabular Processing Test
# ------------------------------------------------------------------------------
def test_process_csv_document(client: TestClient) -> None:
    """Verify uploading and processing a .csv file extracts structured rows and columns."""
    project_id = create_sample_project(client)
    csv_bytes = b"trial_id,accuracy,f1_score\n1,0.94,0.92\n2,0.96,0.95"

    upload_res = client.post(
        f"/api/v1/projects/{project_id}/documents",
        files={"file": ("trials.csv", io.BytesIO(csv_bytes), "text/csv")},
    )
    doc_id = upload_res.json()["id"]

    proc_res = client.post(f"/api/v1/projects/{project_id}/documents/{doc_id}/process")
    assert proc_res.status_code == 200
    assert proc_res.json()["status"] == "extracted"

    content_res = client.get(f"/api/v1/projects/{project_id}/documents/{doc_id}/content")
    assert content_res.status_code == 200
    assert "Columns: trial_id | accuracy | f1_score" in content_res.json()["extracted_text"]
    assert "Row 1: 1 | 0.94 | 0.92" in content_res.json()["extracted_text"]


# ------------------------------------------------------------------------------
# 3. JSON Hierarchical Processing Test
# ------------------------------------------------------------------------------
def test_process_json_document(client: TestClient) -> None:
    """Verify uploading and processing a .json dataset extracts hierarchical text."""
    project_id = create_sample_project(client)
    json_bytes = json.dumps({"experiment": "Transformer_v1", "batch_size": 64}).encode("utf-8")

    upload_res = client.post(
        f"/api/v1/projects/{project_id}/documents",
        files={"file": ("config.json", io.BytesIO(json_bytes), "application/json")},
    )
    doc_id = upload_res.json()["id"]

    proc_res = client.post(f"/api/v1/projects/{project_id}/documents/{doc_id}/process")
    assert proc_res.status_code == 200
    assert proc_res.json()["status"] == "extracted"

    content_res = client.get(f"/api/v1/projects/{project_id}/documents/{doc_id}/content")
    assert "experiment: Transformer_v1" in content_res.json()["extracted_text"]
    assert "batch_size: 64" in content_res.json()["extracted_text"]


# ------------------------------------------------------------------------------
# 4. DOCX Processing Test
# ------------------------------------------------------------------------------
def test_process_docx_document(client: TestClient) -> None:
    """Verify uploading and processing a Word document extracts paragraphs and tables."""
    project_id = create_sample_project(client)

    docx_buf = io.BytesIO()
    doc_obj = docx.Document()
    doc_obj.add_heading("Methodology", level=2)
    doc_obj.add_paragraph("We evaluated linear attention across long contexts.")
    doc_obj.save(docx_buf)
    docx_buf.seek(0)

    upload_res = client.post(
        f"/api/v1/projects/{project_id}/documents",
        files={"file": ("methods.docx", docx_buf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    doc_id = upload_res.json()["id"]

    proc_res = client.post(f"/api/v1/projects/{project_id}/documents/{doc_id}/process")
    assert proc_res.status_code == 200
    assert proc_res.json()["status"] == "extracted"

    content_res = client.get(f"/api/v1/projects/{project_id}/documents/{doc_id}/content")
    assert "## Methodology" in content_res.json()["extracted_text"]
    assert "We evaluated linear attention" in content_res.json()["extracted_text"]


# ------------------------------------------------------------------------------
# 5. Excel Processing Test
# ------------------------------------------------------------------------------
def test_process_excel_document(client: TestClient) -> None:
    """Verify uploading and processing an Excel spreadsheet extracts worksheets and cells."""
    project_id = create_sample_project(client)

    xlsx_buf = io.BytesIO()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Results"
    ws.append(["Model", "Score"])
    ws.append(["GPT-4", "98.2"])
    wb.save(xlsx_buf)
    xlsx_buf.seek(0)

    upload_res = client.post(
        f"/api/v1/projects/{project_id}/documents",
        files={"file": ("results.xlsx", xlsx_buf, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
    )
    doc_id = upload_res.json()["id"]

    proc_res = client.post(f"/api/v1/projects/{project_id}/documents/{doc_id}/process")
    assert proc_res.status_code == 200
    assert proc_res.json()["status"] == "extracted"

    content_res = client.get(f"/api/v1/projects/{project_id}/documents/{doc_id}/content")
    assert "Sheet: Results" in content_res.json()["extracted_text"]
    assert "Model | Score" in content_res.json()["extracted_text"]
    assert "GPT-4 | 98.2" in content_res.json()["extracted_text"]


# ------------------------------------------------------------------------------
# 6. PDF Processing Test
# ------------------------------------------------------------------------------
def test_process_pdf_document(client: TestClient) -> None:
    """Verify uploading and processing a text PDF extracts page text."""
    project_id = create_sample_project(client)
    sample_pdf_bytes = (
        b"%PDF-1.4\n"
        b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n"
        b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n"
        b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >> endobj\n"
        b"4 0 obj << /Length 53 >> stream\n"
        b"BT /F1 24 Tf 100 700 Td (Attention Is All You Need Paper) Tj ET\n"
        b"endstream endobj\n"
        b"5 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n"
        b"xref\n0 6\n0000000000 65535 f \n0000000009 00000 n \n0000000058 00000 n \n0000000115 00000 n \n0000000234 00000 n \n0000000337 00000 n \n"
        b"trailer << /Size 6 /Root 1 0 R >>\nstartxref\n416\n%%EOF\n"
    )

    upload_res = client.post(
        f"/api/v1/projects/{project_id}/documents",
        files={"file": ("attention.pdf", io.BytesIO(sample_pdf_bytes), "application/pdf")},
    )
    doc_id = upload_res.json()["id"]

    proc_res = client.post(f"/api/v1/projects/{project_id}/documents/{doc_id}/process")
    assert proc_res.status_code == 200
    assert proc_res.json()["status"] == "extracted"

    content_res = client.get(f"/api/v1/projects/{project_id}/documents/{doc_id}/content")
    assert "--- Page 1 ---" in content_res.json()["extracted_text"]
    assert "Attention Is All You Need" in content_res.json()["extracted_text"]


# ------------------------------------------------------------------------------
# 7. Error Handling & Reprocessing Tests
# ------------------------------------------------------------------------------
def test_process_empty_document_marks_failed(client: TestClient) -> None:
    """Verify processing an empty file marks status as 'failed' and captures processing_error."""
    project_id = create_sample_project(client)
    upload_res = client.post(
        f"/api/v1/projects/{project_id}/documents",
        files={"file": ("blank.txt", io.BytesIO(b"   \n\n   "), "text/plain")},
    )
    doc_id = upload_res.json()["id"]

    proc_res = client.post(f"/api/v1/projects/{project_id}/documents/{doc_id}/process")
    assert proc_res.status_code == 200
    data = proc_res.json()
    assert data["status"] == "failed"
    assert data["processing_error"] is not None
    assert "blank or unreadable" in data["processing_error"]

    content_res = client.get(f"/api/v1/projects/{project_id}/documents/{doc_id}/content")
    assert content_res.status_code == 200
    assert content_res.json()["status"] == "failed"
    assert content_res.json()["extracted_text"] is None


def test_process_nonexistent_document_returns_404(client: TestClient) -> None:
    """Verify processing a non-existent document ID returns 404 Not Found."""
    project_id = create_sample_project(client)
    random_doc_id = str(uuid.uuid4())
    res = client.post(f"/api/v1/projects/{project_id}/documents/{random_doc_id}/process")
    assert res.status_code == 404


def test_process_document_cross_project_isolation(client: TestClient) -> None:
    """Verify attempting to process project A's document via project B's URL returns 404."""
    project_a = create_sample_project(client)
    project_b = create_sample_project(client)

    upload_res = client.post(
        f"/api/v1/projects/{project_a}/documents",
        files={"file": ("project_a_doc.txt", io.BytesIO(b"Valid content"), "text/plain")},
    )
    doc_id = upload_res.json()["id"]

    # Try processing via project B
    res = client.post(f"/api/v1/projects/{project_b}/documents/{doc_id}/process")
    assert res.status_code == 404
