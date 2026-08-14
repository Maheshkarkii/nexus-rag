"""DOCX document extractor for Word files, paragraphs, headings, and tables using python-docx."""

import logging
from pathlib import Path
from typing import List
import docx
from app.db.models.document import Document
from app.services.document_processing.base import BaseDocumentProcessor, ExtractionResult

logger = logging.getLogger("ai_research_assistant.processors.docx")


class DocxProcessor(BaseDocumentProcessor):
    """Extracts structured text from Microsoft Word (.docx) files including paragraphs and tables."""

    async def extract(self, file_path: Path, document: Document) -> ExtractionResult:
        if not file_path.exists():
            raise FileNotFoundError(f"Physical DOCX file not found at: {file_path}")

        try:
            doc = docx.Document(str(file_path))
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX document structure: {e}") from e

        content_blocks: List[str] = []
        paragraph_count = 0
        table_count = 0

        # 1. Extract paragraphs and headings
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraph_count += 1
                style_name = para.style.name if para.style else "Normal"
                if "heading" in style_name.lower():
                    content_blocks.append(f"## {text}")
                else:
                    content_blocks.append(text)

        # 2. Extract tables in readable markdown-style tabular format
        for t_idx, table in enumerate(doc.tables, start=1):
            table_count += 1
            table_rows: List[str] = []
            for row in table.rows:
                row_cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                # Filter out redundant merged cell text repetitions
                unique_cells = []
                for cell in row_cells:
                    if not unique_cells or cell != unique_cells[-1]:
                        unique_cells.append(cell)
                if any(unique_cells):
                    table_rows.append(" | ".join(unique_cells))

            if table_rows:
                content_blocks.append(f"[Table {t_idx}]\n" + "\n".join(table_rows))

        combined_text = "\n\n".join(content_blocks).strip()

        if not combined_text:
            raise ValueError("DOCX document contains no readable text, paragraphs, or table entries.")

        char_count = len(combined_text)
        word_count = len(combined_text.split())

        return ExtractionResult(
            text=combined_text,
            character_count=char_count,
            word_count=word_count,
            metadata={
                "paragraph_count": paragraph_count,
                "table_count": table_count,
            },
        )
