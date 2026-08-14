import io
import re
from typing import Any

import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.pdfgen import canvas
from reportlab.platypus import (
    HRFlowable,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


class NumberedCanvas(canvas.Canvas):
    """Two-pass canvas for adding page numbers in 'Page X of Y' format."""

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_number(num_pages)
            super().showPage()
        super().save()

    def draw_page_number(self, page_count):
        self.saveState()
        self.setFont("Helvetica", 9)
        self.setFillColor(colors.HexColor("#64748b"))
        page_text = f"Page {self._pageNumber} of {page_count}"
        self.drawRightString(612 - 54, 36, page_text)
        self.drawString(54, 36, "AI Research Assistant — Confidential Report")
        self.restoreState()


class MarkdownExporter:
    """Exports structured report JSON into standard Markdown format."""

    @staticmethod
    def export(content_json: dict[str, Any]) -> str:
        title = content_json.get("title", "Research Report")
        report_type = content_json.get("report_type", "research_summary").replace("_", " ").title()
        sections = content_json.get("sections", [])
        sources = content_json.get("sources", [])

        md = []
        md.append(f"# {title}\n")
        md.append(f"**Report Type:** {report_type}\n")
        md.append("---\n")

        for section in sections:
            sec_title = section.get("title", "Untitled Section")
            sec_content = section.get("content", "")
            md.append(f"## {sec_title}\n")
            md.append(f"{sec_content}\n")

        if sources:
            md.append("## Sources\n")
            for idx, src in enumerate(sources, 1):
                sid = src.get("source_id", f"S{idx}")
                fname = src.get("filename", "Unknown Document")
                loc = src.get("location_info", "")
                loc_str = f" ({loc})" if loc else ""
                md.append(f"{idx}. **[{sid}]** {fname}{loc_str}")
            md.append("")

        return "\n".join(md)


class PDFExporter:
    """Exports structured report JSON into a formatted PDF document."""

    @staticmethod
    def export(content_json: dict[str, Any]) -> bytes:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=54,
            leftMargin=54,
            topMargin=54,
            bottomMargin=54,
        )

        styles = getSampleStyleSheet()
        
        # Custom typography styles
        title_style = ParagraphStyle(
            "DocTitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=22,
            leading=26,
            textColor=colors.HexColor("#0f172a"),
            alignment=0,
            spaceAfter=6,
        )

        meta_style = ParagraphStyle(
            "DocMeta",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=10,
            leading=14,
            textColor=colors.HexColor("#64748b"),
            spaceAfter=14,
        )

        h2_style = ParagraphStyle(
            "Heading2Custom",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=14,
            leading=18,
            textColor=colors.HexColor("#1e293b"),
            spaceBefore=14,
            spaceAfter=8,
        )

        body_style = ParagraphStyle(
            "BodyCustom",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10,
            leading=14,
            textColor=colors.HexColor("#334155"),
            spaceAfter=10,
        )

        story = []

        # Title & Metadata
        title = content_json.get("title", "Research Report")
        report_type = content_json.get("report_type", "research_summary").replace("_", " ").title()
        
        story.append(Paragraph(title, title_style))
        story.append(Paragraph(f"<b>Type:</b> {report_type}", meta_style))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#e2e8f0"), spaceAfter=14))

        # Sections
        sections = content_json.get("sections", [])
        for section in sections:
            sec_title = section.get("title", "Untitled Section")
            sec_content = section.get("content", "")
            
            story.append(Paragraph(sec_title, h2_style))
            
            # Clean markdown paragraphs and list items
            paras = sec_content.split("\n\n")
            for p in paras:
                p_clean = p.replace("\n", " ").strip()
                if p_clean:
                    # Convert markdown bold to PDF HTML tags
                    p_formatted = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", p_clean)
                    story.append(Paragraph(p_formatted, body_style))

            story.append(Spacer(1, 6))

        # Sources
        sources = content_json.get("sources", [])
        if sources:
            story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#e2e8f0"), spaceBefore=12, spaceAfter=10))
            story.append(Paragraph("Sources & Citations", h2_style))
            
            table_data = [["ID", "Document", "Location Details"]]
            for idx, src in enumerate(sources, 1):
                sid = src.get("source_id", f"S{idx}")
                fname = src.get("filename", "Unknown Document")
                loc = src.get("location_info", "-")
                table_data.append([sid, fname, loc])

            src_table = Table(table_data, colWidths=[50, 250, 204])
            src_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f1f5f9")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0f172a")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
            ]))
            story.append(src_table)

        doc.build(story, canvasmaker=NumberedCanvas)
        return buffer.getvalue()


class DOCXExporter:
    """Exports structured report JSON into a formatted Word (.docx) document."""

    @staticmethod
    def export(content_json: dict[str, Any]) -> bytes:
        doc = docx.Document()

        # Set page margins
        for sec in doc.sections:
            sec.top_margin = Inches(0.75)
            sec.bottom_margin = Inches(0.75)
            sec.left_margin = Inches(0.75)
            sec.right_margin = Inches(0.75)

        title = content_json.get("title", "Research Report")
        report_type = content_json.get("report_type", "research_summary").replace("_", " ").title()

        # Title
        p_title = doc.add_paragraph()
        run_title = p_title.add_run(title)
        run_title.font.name = "Calibri"
        run_title.font.size = Pt(22)
        run_title.font.bold = True
        run_title.font.color.rgb = RGBColor(15, 23, 42)

        # Meta
        p_meta = doc.add_paragraph()
        run_meta = p_meta.add_run(f"Report Type: {report_type}")
        run_meta.font.name = "Calibri"
        run_meta.font.size = Pt(10)
        run_meta.font.color.rgb = RGBColor(100, 116, 139)

        doc.add_paragraph("―" * 55)

        # Sections
        sections = content_json.get("sections", [])
        for section in sections:
            sec_title = section.get("title", "Untitled Section")
            sec_content = section.get("content", "")

            # Heading
            h = doc.add_heading(sec_title, level=2)
            h.style.font.color.rgb = RGBColor(30, 41, 59)

            # Paragraphs
            paras = sec_content.split("\n\n")
            for p_text in paras:
                p_clean = p_text.replace("\n", " ").strip()
                if p_clean:
                    doc.add_paragraph(p_clean)

        # Sources
        sources = content_json.get("sources", [])
        if sources:
            doc.add_heading("Sources & Citations", level=2)
            table = doc.add_table(rows=1, cols=3)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "ID"
            hdr_cells[1].text = "Document"
            hdr_cells[2].text = "Location Details"

            for idx, src in enumerate(sources, 1):
                row_cells = table.add_row().cells
                row_cells[0].text = src.get("source_id", f"S{idx}")
                row_cells[1].text = src.get("filename", "Unknown Document")
                row_cells[2].text = src.get("location_info", "-")

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
